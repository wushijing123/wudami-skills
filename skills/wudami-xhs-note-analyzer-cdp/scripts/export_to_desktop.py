#!/usr/bin/env python3
import os
import sys
import json
import shutil
import subprocess
from datetime import datetime
import argparse

def ensure_dependencies():
    missing = []
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
        
    try:
        import markdown
    except ImportError:
        missing.append("markdown")
        
    if missing:
        print(f"📦 正在自动安装导出模块: {missing}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", *missing, "-q"])

def md_to_docx(md_path, docx_path):
    ensure_dependencies()
    import docx
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote'
        else:
            doc.add_paragraph(line)
            
    try:
        doc.save(docx_path)
    except Exception as e:
        print(f"⚠️ 保存 DOCX 失败: {e}")

def inject_frames_into_md(lines, frames_dir):
    import re
    import base64
    
    new_lines = []
    frame_files = sorted([f for f in os.listdir(frames_dir) if f.startswith('frame_')]) if os.path.exists(frames_dir) else []
    
    def get_base64_img(frame_idx):
        if frame_idx < 1: frame_idx = 1
        if frame_idx > len(frame_files): frame_idx = len(frame_files)
        if not frame_files: return ""
        
        frame_file = frame_files[frame_idx - 1]
        frame_path = os.path.join(frames_dir, frame_file)
        if not os.path.exists(frame_path): return ""
        
        with open(frame_path, "rb") as img_file:
            b64 = base64.b64encode(img_file.read()).decode('utf-8')
        return f'<div style="position: relative; display: inline-block;" onmouseover="this.querySelector(\'.popup\').style.display=\'block\'" onmouseout="this.querySelector(\'.popup\').style.display=\'none\'"><img src="data:image/jpeg;base64,{b64}" width="160" style="border-radius: 6px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); cursor: zoom-in; transition: transform 0.2s;" onmouseenter="this.style.transform=\'scale(1.05)\'" onmouseleave="this.style.transform=\'scale(1)\'" /><img class="popup" src="data:image/jpeg;base64,{b64}" style="display: none; position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%); z-index: 9999; max-height: 90vh; max-width: 90vw; border-radius: 12px; box-shadow: 0 20px 60px rgba(0,0,0,0.6); pointer-events: none;" /></div>'

    in_target_table = False
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|'):
            parts = line.split('|')
            # The first real column string
            col1 = parts[1].strip() if len(parts) > 1 else ""
            
            # Detect target table (兼容视频的"时间/镜头"与图文的"页面序号/排版")
            if ('时间' in col1 or '页面序号' in col1 or '序号' in col1) and ('镜头' in line or '排版' in line or '视觉' in line or '拆解' in line) and not '封面' in col1 and not '图' in col1:
                in_target_table = True
                if '对应画面' not in line:
                    parts.insert(2, ' 对应画面 ')
                    line = '|'.join(parts)
            
            # Detect separator row
            elif in_target_table and '---' in col1:
                parts.insert(2, ' --- ')
                line = '|'.join(parts)
                
            # Detect data rows
            elif in_target_table:
                time_match = re.search(r'(\d+):(\d+)', col1)
                img_match = re.search(r'(图|p|P)\s*(\d+)', col1, re.IGNORECASE)
                cover_match = '封面' in col1 or '首图' in col1
                tail_match = '尾图' in col1 or '最后一图' in col1

                frame_idx = None
                if time_match:
                    start_sec = int(time_match.group(1)) * 60 + int(time_match.group(2))
                    frame_idx = (start_sec // 3) + 1  # 3秒一帧
                elif img_match:
                    frame_idx = int(img_match.group(2))
                    # 由于小红书图文笔记的展示习惯，如果出现了“封面”占据了第1张图，
                    # 那么后续的“图1”或者“p1”一般指的是内页的第1张（即文件系统里第2张图片）。
                    if '内页' in col1 or img_match.group(1).lower() == 'p' or img_match.group(1) == '图':
                        frame_idx += 1
                elif cover_match:
                    frame_idx = 1
                elif tail_match:
                    if os.path.exists(frames_dir):
                        frame_files_list = [f for f in os.listdir(frames_dir) if f.startswith('frame_') and f.endswith('.jpg')]
                        if frame_files_list:
                            frame_idx = len(frame_files_list)
                    
                img_tag = get_base64_img(frame_idx) if frame_idx is not None else ""
                parts.insert(2, f' {img_tag} ')
                line = '|'.join(parts)
                
        else:
            in_target_table = False
            
        new_lines.append(line.rstrip('\n') + '\n')
    return new_lines

def md_to_html_with_frames(md_path, html_path, json_path):
    ensure_dependencies()
    import markdown
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    frames_dir = os.path.join(os.path.dirname(json_path), "_frames")
    lines = inject_frames_into_md(lines, frames_dir)
    md_text = "".join(lines)
    
    html_content = markdown.markdown(md_text, extensions=['tables'])
    
    full_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>小红书深度拆解战报</title>
    <style>
        :root {{ --primary: #ff2442; --bg: #f9f9f9; --text: #333; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", sans-serif; line-height: 1.6; color: var(--text); background: var(--bg); margin: 0; padding: 40px 20px; }}
        .container {{ max-width: 900px; margin: 0 auto; background: #fff; padding: 40px 60px; border-radius: 16px; box-shadow: 0 10px 30px rgba(0,0,0,0.05); }}
        h1 {{ color: var(--primary); border-bottom: 2px solid #ffeeef; padding-bottom: 12px; margin-top: 0; text-align: center; font-size: 28px; }}
        h2 {{ color: #1a1a1a; margin-top: 32px; font-size: 22px; position: relative; padding-left: 16px; }}
        h2::before {{ content: ''; position: absolute; left: 0; top: 4px; bottom: 4px; width: 4px; background: var(--primary); border-radius: 2px; }}
        h3 {{ color: var(--primary); font-size: 18px; margin-top: 24px; margin-bottom: 16px; padding: 8px 16px; background: #fff5f5; border-left: 4px solid var(--primary); border-radius: 0 6px 6px 0; font-weight: bold; }}
        h4 {{ color: #333; font-size: 16px; margin-top: 20px; margin-bottom: 12px; padding-left: 12px; border-left: 3px solid #ff99a8; font-weight: 600; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 15px; border-radius: 8px; overflow: hidden; box-shadow: 0 0 0 1px #eee; }}
        th, td {{ padding: 16px; text-align: left; vertical-align: top; border-bottom: 1px solid #eee; }}
        th {{ background-color: #fafafa; font-weight: 600; color: #555; }}
        tr:last-child td {{ border-bottom: none; }}
        tr:hover {{ background-color: #fdfdfd; }}
        blockquote {{ border-left: 4px solid var(--primary); padding: 12px 20px; color: #666; margin: 20px 0; background: #fff8f8; border-radius: 0 8px 8px 0; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 4px; font-family: monospace; color: #d63384; }}
    </style>
</head>
<body>
    <div class="container">
        {html_content}
    </div>
</body>
</html>
"""
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(full_html)
        
def export_to_desktop(md_path, html_path, json_path):
    nickname = "未知账号"
    is_note = False
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                if 'account' in data and 'nickname' in data['account']:
                    nickname = data['account']['nickname'].strip()
                elif 'nickname' in data:
                    nickname = data['nickname'].strip()
                elif 'note' in data and 'title' in data['note']:
                    nickname = data['note']['title'].strip()
                    is_note = True
                elif 'title' in data:
                    nickname = data['title'].strip()
                    is_note = True
            except Exception:
                pass

    # Sanitize nickname (remove invalid filename chars)
    nickname = "".join(x for x in nickname if x.isalnum() or x in "._- 　【】()（）")
    # Limit length to avoid path too long errors
    if len(nickname) > 40:
        nickname = nickname[:40] + "..."

    date_str = datetime.now().strftime("%Y-%m-%d")
    desktop_dir = os.path.expanduser("~/Desktop")
    
    # 为每篇笔记创建专属的分类子文件夹，避免根目录太乱
    obsidian_base_dir = os.path.expanduser("~/Obsidian仓库/吴大咪一人公司/02-素材库/对标账号库/笔记拆解档案")
    folder_name = f"{date_str}-{nickname}"
    obsidian_dir = os.path.join(obsidian_base_dir, folder_name)
    has_obsidian = os.path.exists(obsidian_base_dir)
    
    if has_obsidian:
        os.makedirs(obsidian_dir, exist_ok=True)
    
    # 桌面端导出的文件保持长前缀，因为桌面没有文件夹分类
    html_target = os.path.join(desktop_dir, f"{date_str}-{nickname}-可视化战报.html")
    docx_target = os.path.join(desktop_dir, f"{date_str}-{nickname}-深度拆解报告.docx")
    
    # 知识库端文件存入子文件夹，简化文件名
    md_target = os.path.join(obsidian_dir, "拆解报告.md") if has_obsidian else None
    
    
    # 转换为精美带截图的 HTML 并推送桌面
    if os.path.exists(md_path):
        md_to_html_with_frames(md_path, html_target, json_path)
        print(f"✅ HTML 视觉内嵌战报已推送至桌面: {html_target}")
        
        # 将 HTML 也同步一份到知识库方便查看
        if has_obsidian:
            html_obsidian = os.path.join(obsidian_dir, "可视化战报.html")
            try:
                shutil.copy(html_target, html_obsidian)
                print(f"✅ HTML 可视化战报已同步至知识库: {html_obsidian}")
            except Exception as e:
                print(f"⚠️ HTML 同步至 Obsidian 失败: {e}")
        
    # 转换为 DOCX
    if os.path.exists(md_path):
        # 自动转换为 DOCX 并推送桌面，方便给没有用 Obsidian 的团队成员或客户查看
        md_to_docx(md_path, docx_target)
        print(f"✅ DOCX 深度拆解报告已推送至桌面: {docx_target}")
        
        # 如果存在指定的Obsidian目录，同步一份Markdown到Obsidian
        if has_obsidian:
            try:
                shutil.copy(md_path, md_target)
                print(f"✅ Markdown 源文件已同步至知识库: {md_target}")
            except Exception as e:
                print(f"⚠️ 同步到 Obsidian 失败: {e}")
    else:
        print(f"⚠️ Markdown 报告未找到: {md_path}")

    # 自动同步故事版（storyboard.jpg）到 Obsidian
    storyboard_src = os.path.join(os.path.dirname(json_path), "storyboard.jpg")
    if os.path.exists(storyboard_src) and has_obsidian:
        storyboard_target = os.path.join(obsidian_dir, "storyboard.jpg")
        try:
            shutil.copy(storyboard_src, storyboard_target)
            print(f"✅ 视频故事版已同步至知识库: {storyboard_target}")
        except Exception as e:
            print(f"⚠️ 故事版同步失败: {e}")
        
    print("\n💎 所有拆解资产已成功打包到您的桌面，并同步至黑曜石知识库！")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="一键推送报告到桌面")
    parser.add_argument("--md", required=True, help="Markdown报告路径")
    parser.add_argument("--html", required=True, help="HTML报告路径")
    parser.add_argument("--json", required=True, help="数据JSON路径")
    args = parser.parse_args()
    
    export_to_desktop(args.md, args.html, args.json)
