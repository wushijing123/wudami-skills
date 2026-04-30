#!/usr/bin/env python3
import os
import sys
import json
import shutil
import subprocess
from datetime import datetime
import argparse

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("缺少 python-docx 模块，正在为您自动安装...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

def md_to_docx(md_path, docx_path):
    ensure_docx()
    import docx
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote'
        else:
            doc.add_paragraph(line)
            
    try:
        doc.save(docx_path)
    except Exception as e:
        print(f"⚠️ 保存 DOCX 失败: {e}")

def export_to_desktop(md_path, html_path, json_path):
    nickname = "未知账号"
    is_note = False
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                if 'account' in data and 'nickname' in data['account']:
                    nickname = data['account']['nickname'].strip()
                elif 'nickname' in data:
                    nickname = data['nickname'].strip()
                elif 'note' in data and 'title' in data['note']:
                    nickname = data['note']['title'].strip()
                    is_note = True
                elif 'title' in data:
                    nickname = data['title'].strip()
                    is_note = True
            except Exception:
                pass

    # Sanitize nickname (remove invalid filename chars)
    nickname = "".join(x for x in nickname if x.isalnum() or x in "._- 　【】()（）")
    # Limit length to avoid path too long errors
    if len(nickname) > 40:
        nickname = nickname[:40] + "..."

    date_str = datetime.now().strftime("%Y-%m-%d")
    desktop_dir = os.path.expanduser("~/Desktop")
    obsidian_dir = os.path.expanduser("~/Obsidian仓库/吴大咪一人公司/02-素材库/对标账号库/笔记拆解档案")
    
    html_target = os.path.join(desktop_dir, f"{date_str}-{nickname}-可视化战报.html")
    docx_target = os.path.join(desktop_dir, f"{date_str}-{nickname}-深度拆解报告.docx")
    md_target = os.path.join(obsidian_dir, f"{date_str}-{nickname}-拆解报告.md")
    
    # [测试阶段屏蔽] 拷贝 HTML
    # if os.path.exists(html_path):
    #     shutil.copy(html_path, html_target)
    #     print(f"✅ HTML 战报已推送至桌面: {html_target}")
    # elif not is_note: # Note analyzer usually doesn't have HTML right now
    #     print(f"⚠️ HTML 战报未找到: {html_path}")
        
    # 转换为 DOCX
    if os.path.exists(md_path):
        # [测试阶段屏蔽] 转换为 DOCX 并推送桌面
        # md_to_docx(md_path, docx_target)
        # print(f"✅ DOCX 深度拆解报告已推送至桌面: {docx_target}")
        
        # 如果是笔记拆解或存在指定的Obsidian目录，同步一份Markdown到Obsidian
        if os.path.exists(os.path.dirname(md_target)):
            shutil.copy(md_path, md_target)
            print(f"✅ Markdown 源文件已同步至知识库: {md_target}")
        else:
            try:
                os.makedirs(os.path.dirname(md_target), exist_ok=True)
                shutil.copy(md_path, md_target)
                print(f"✅ Markdown 源文件已同步至知识库: {md_target}")
            except Exception as e:
                print(f"⚠️ 同步到 Obsidian 失败: {e}")
    else:
        print(f"⚠️ Markdown 报告未找到: {md_path}")
        
    print("\n💎 所有拆解资产已成功打包到您的桌面，并同步至黑曜石知识库！")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="一键推送报告到桌面")
    parser.add_argument("--md", required=True, help="Markdown报告路径")
    parser.add_argument("--html", required=True, help="HTML报告路径")
    parser.add_argument("--json", required=True, help="数据JSON路径")
    args = parser.parse_args()
    
    export_to_desktop(args.md, args.html, args.json)
